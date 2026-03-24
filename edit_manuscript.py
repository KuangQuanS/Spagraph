"""
Edit manuscript.docx:
1. Delete formula #4 (para 19) and its intro text (para 18), merge para 20 text
2. Delete formula #5 (para 23) and rewrite para 22 ending
3. Modify formula #6 (para 26) from LeakyReLU to MLP
4. Delete formula #8 (para 30) and its intro (para 29), merge with para 31
5. Delete formula #9 (para 32) and "with lres=0.5" (para 33), merge with para 31
6. Delete degree-scaling description (para 59) - rewrite without degree correction
7. Fix ARS formula #22 (para 71) - add N to denominator
8. Renumber all formulas after deletions
"""
import sys, io, copy, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from lxml import etree
from docx import Document
from docx.oxml.ns import qn

doc = Document('D:/Spagraph/manuscript.docx')

ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
      'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}

body = doc.element.body
paragraphs = body.findall(qn('w:p'))

def get_text(p):
    return ''.join(p.itertext()).strip()

def remove_paragraph(p):
    parent = p.getparent()
    parent.remove(p)

# ============================================================
# 1. Delete formula #4 (para 19) and its intro (para 18)
# ============================================================
print("1. Deleting formula #4 and intro...")
p18 = paragraphs[18]
p19 = paragraphs[19]

# Clear para 18 and rewrite
for child in list(p18):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag != 'pPr':
        p18.remove(child)

new_run = etree.SubElement(p18, qn('w:r'))
rpr = etree.SubElement(new_run, qn('w:rPr'))
rfont = etree.SubElement(rpr, qn('w:rFonts'))
rfont.set(qn('w:ascii'), 'Times New Roman')
rfont.set(qn('w:hAnsi'), 'Times New Roman')
t = etree.SubElement(new_run, qn('w:t'))
t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
t.text = 'The total loss function for the alignment stage combines the ELBO losses from both modalities with an MMD alignment term, weighted by \u03bbmmd (set to 0.1).'

remove_paragraph(p19)

# Fix para 20: remove "where lmmd ... constraint." prefix
p20 = paragraphs[20]
# Remove all children except pPr, then rewrite
for child in list(p20):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag != 'pPr':
        p20.remove(child)

new_run = etree.SubElement(p20, qn('w:r'))
rpr = etree.SubElement(new_run, qn('w:rPr'))
rfont = etree.SubElement(rpr, qn('w:rFonts'))
rfont.set(qn('w:ascii'), 'Times New Roman')
rfont.set(qn('w:hAnsi'), 'Times New Roman')
t = etree.SubElement(new_run, qn('w:t'))
t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
t.text = 'Optimization was performed using the Adam algorithm with a cosine annealing learning rate schedule. Upon model convergence, we generated a robust \u201cprototype embedding\u201d for each cell type cluster by averaging the latent mean vectors of all cells assigned to that cluster in the scRNA-seq reference. These prototypes serve as noise-robust semantic anchors for the subsequent spatial deconvolution.'

print(f"  Para 18 now: {get_text(p18)[:100]}")
print(f"  Para 20 now: {get_text(p20)[:100]}")

# ============================================================
# 2. Delete formula #5 (para 23) and fix para 22 ending
# ============================================================
print("\n2. Deleting formula #5...")
p22 = paragraphs[22]
p23 = paragraphs[23]

# Fix para 22 ending: change colon to period and add description
# Find the last text run
for child in reversed(list(p22)):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'r':
        t_elem = child.find(qn('w:t'))
        if t_elem is not None and t_elem.text and t_elem.text.strip():
            if t_elem.text.rstrip().endswith(':'):
                t_elem.text = t_elem.text.rstrip()[:-1] + ', denoted as h and g respectively.'
            break

remove_paragraph(p23)
print(f"  Formula #5 deleted, para 22 ending fixed")

# ============================================================
# 3. Modify formula #6 (para 26): LeakyReLU -> MLP
# ============================================================
print("\n3. Modifying formula #6...")
p25 = paragraphs[25]
p26 = paragraphs[26]
p27 = paragraphs[27]

# Change LeakyReLU to MLP in formula
for mt in p26.iter(qn('m:t')):
    if mt.text:
        if 'LeakyReLU' in mt.text:
            mt.text = mt.text.replace('LeakyReLU', 'MLP')
            print(f"  Changed LeakyReLU -> MLP")

# Change 'a' (the attention vector) to theta in the formula
# Need to be careful - find the specific 'a' that's the attention parameter
# In the formula: a^T[Wh_i || Wh_j] -> MLP_theta(h_i || h_j)
for mt in p26.iter(qn('m:t')):
    if mt.text and mt.text.strip() == 'a':
        mt.text = '\u03b8'
        print(f"  Changed 'a' -> '\u03b8'")
        break

# Remove the T superscript (transpose) - find it
for mt in p26.iter(qn('m:t')):
    if mt.text and mt.text.strip() == 'T':
        mt.text = ''
        print(f"  Removed 'T' superscript")
        break

# Update para 27: "shared linear transformation W" -> "scoring MLP"
for child in list(p27):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'r':
        t_elem = child.find(qn('w:t'))
        if t_elem is not None and t_elem.text:
            if 'shared linear transformation' in t_elem.text:
                t_elem.text = t_elem.text.replace(
                    'shared linear transformation',
                    'scoring MLP'
                )
                print(f"  Updated para 27 text")
            if 'enabling attention-based message passing across heterogeneous node types' in t_elem.text:
                t_elem.text = t_elem.text.replace(
                    'enabling attention-based message passing across heterogeneous node types',
                    'which jointly processes concatenated node features to compute attention scores across heterogeneous node types'
                )

# Update inline math W -> theta in para 27
for omath in p27.findall(qn('m:oMath')):
    mt_text = ''.join(omath.itertext()).strip()
    if mt_text == 'W':
        for m_t in omath.iter(qn('m:t')):
            if m_t.text == 'W':
                m_t.text = '\u03b8'
                print(f"  Changed inline W -> \u03b8 in para 27")

# ============================================================
# 4-5. Delete formulas #8, #9 and merge text
# ============================================================
print("\n4-5. Deleting formulas #8, #9...")
p29 = paragraphs[29]
p30 = paragraphs[30]
p31 = paragraphs[31]
p32 = paragraphs[32]
p33 = paragraphs[33]

# Rewrite para 29 to combine the content of #8 and #9
for child in list(p29):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag != 'pPr':
        p29.remove(child)

new_run = etree.SubElement(p29, qn('w:r'))
rpr = etree.SubElement(new_run, qn('w:rPr'))
rfont = etree.SubElement(rpr, qn('w:rFonts'))
rfont.set(qn('w:ascii'), 'Times New Roman')
rfont.set(qn('w:hAnsi'), 'Times New Roman')
t = etree.SubElement(new_run, qn('w:t'))
t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
t.text = 'The updated spot representation is computed as a weighted sum of transformed neighbor features. To mitigate the over-smoothing phenomenon common in deep GNNs, where node representations become indistinguishable, we incorporated a residual fusion mechanism with \u03bbres = 0.5, balancing the GAT-aggregated feature with the initial latent representation to preserve the original signal identity.'

remove_paragraph(p30)
remove_paragraph(p31)
remove_paragraph(p32)
remove_paragraph(p33)
print(f"  Merged and deleted paras 30-33")

# ============================================================
# 6. Remove degree-scaling correction (para 59)
# ============================================================
print("\n6. Removing degree-scaling correction...")
p59 = paragraphs[59]

for child in list(p59):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag != 'pPr':
        p59.remove(child)

new_run = etree.SubElement(p59, qn('w:r'))
rpr = etree.SubElement(new_run, qn('w:rPr'))
rfont = etree.SubElement(rpr, qn('w:rFonts'))
rfont.set(qn('w:ascii'), 'Times New Roman')
rfont.set(qn('w:hAnsi'), 'Times New Roman')
t = etree.SubElement(new_run, qn('w:t'))
t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
t.text = 'Following self-supervised pre-training, the learned attention weights serve as the primary indicator of communication significance. We construct a high-fidelity communication topology by applying a robust confidence threshold, retaining only those interactions that significantly exceed the random background expectation.'

print(f"  Degree-scaling removed, para 59 rewritten")

# ============================================================
# 7. Fix ARS formula #22 (para 71) - add N to denominator
# ============================================================
print("\n7. Fixing ARS formula #22...")
p71 = paragraphs[71]

found_4 = False
for mt in p71.iter(qn('m:t')):
    if mt.text and mt.text.strip() == '4':
        mt.text = '4N'
        found_4 = True
        print(f"  Changed '4' to '4N' in ARS denominator")
        break

if not found_4:
    print("  WARNING: Could not find '4' in ARS formula")

# ============================================================
# 8. Renumber formulas
# ============================================================
print("\n8. Renumbering formulas...")

# Original numbers present (after deleting #4, #5, #8, #9):
# 1, 2, 3, 6, 7, 10, 11, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22
# New numbering:
# 1, 2, 3, 4, 5,  6,  7,  8,  9, 10, 11, 12, 13, 14, 15, 16, 17

old_to_new = {
    1: 1, 2: 2, 3: 3,
    6: 4, 7: 5,
    10: 6, 11: 7,
    13: 8, 14: 9, 15: 10,
    16: 11, 17: 12,
    18: 13, 19: 14, 20: 15, 21: 16, 22: 17
}

# Process in reverse order (largest first) to avoid collision
remaining_paragraphs = body.findall(qn('w:p'))
for p in remaining_paragraphs:
    for omath in p.findall('.//' + qn('m:oMath')):
        text = ''.join(omath.itertext())
        if '#' in text:
            for mt in omath.iter(qn('m:t')):
                if mt.text and '#' in mt.text:
                    # Extract the number after #
                    match = re.search(r'#(\d+)', mt.text)
                    if match:
                        old_num = int(match.group(1))
                        if old_num in old_to_new:
                            new_num = old_to_new[old_num]
                            if old_num != new_num:
                                # Use a temporary placeholder to avoid collision
                                mt.text = mt.text.replace(f'#{old_num}', f'#TEMP{new_num}')
                                print(f"  #{old_num} -> #{new_num}")

# Second pass: remove TEMP prefix
remaining_paragraphs = body.findall(qn('w:p'))
for p in remaining_paragraphs:
    for mt in p.iter(qn('m:t')):
        if mt.text and '#TEMP' in mt.text:
            mt.text = mt.text.replace('#TEMP', '#')

# ============================================================
# Save
# ============================================================
print("\nSaving...")
doc.save('D:/Spagraph/manuscript.docx')
print("Done! Saved to D:/Spagraph/manuscript.docx")

# Verify
print("\n=== Verification: All numbered formulas ===")
doc2 = Document('D:/Spagraph/manuscript.docx')
for i, p in enumerate(doc2.paragraphs):
    math_elems = p._element.findall('.//' + qn('m:oMath'))
    for me in math_elems:
        mt = ''.join(me.itertext())
        if '#' in mt:
            print(f"  [{i}] {mt[:80]}")
